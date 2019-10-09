# GenerateReport.py

from datetime import datetime
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import StockData
import CombinePics

# Cleans up existing png files generated for PDF
def cleanUpFiles(tickers, files):

    for ticker in tickers:
        for file in files:
            if os.path.exists(ticker + file):
                os.remove(ticker + file)
            else:
                return False

    return True

# checks to see if two files exist
def fileExists(ticker):

    if os.path.exists(ticker + "DailyPrice.png"):
        if os.path.exists(ticker + "DailyVolume.png"):
            return True

    return False

# Creates a cover page for the document
def createCover(document):

    today = datetime.now()
    datestr = today.strftime("%m-%d-%y")

    p = document.add_paragraph()
    for i in range(11):
        p.add_run('\n')

    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r = p.add_run("Stock Market Report\n")
    r.font.size = Pt(22)
    r.bold = True
    t = p.add_run(datestr)
    t.font.size = Pt(16)
    t.bold = True

    document.add_page_break()

    return

# Function to create a table of contents
def createTOC(document, tickers):

    # Add a page number
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("1")
    r.bold = True

    document.add_heading("Table of Contents", level = 0)

    p = document.add_paragraph()

    pageNum = 2

    for ticker in tickers:
        string = ticker
        for i in range(len(ticker), 75 - len(str(pageNum))):
            string += "."

        s = p.add_run(string + str(pageNum) + "\n")
        pageNum += 1

    document.add_page_break()

    return

# Creates a section for each ticker to display general info,
# plots and other market data
def writeTicker(document, ticker, T, pageNum):

    d = T[ticker][0]
    keys = T[ticker][1]
    values = T[ticker][2]

    # Add a page number
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(str(pageNum))
    r.bold = True

    document.add_heading(ticker, level = 1)

    # set the document style: Times New Roman
    p = document.add_paragraph()
    style = document.styles["Normal"]
    font = style.font
    font.size = Pt(12)
    font.name = "Times New Roman"
    p.style = document.styles["Normal"]

    # Format and write the price change info
    p.add_run("The daily price movement and volume of " + ticker + " are shown below.  ")
    p.add_run(ticker + " opened the day at " + d["Open"] + " and closed the day at " +
              d["Close"] + " for a gain of ")
    s = p.add_run(str("%.3f" % d["Change"]) + "%.  ")

    if d["Change"] >= 0:
        s.font.color.rgb = RGBColor(0, 255, 0)
    else:
        s.font.color.rgb = RGBColor(255, 0, 0)

    # Check if the picture files exist before trying to insert them
    if fileExists(ticker):
        CombinePics.combine(ticker + "DailyVolume.png", ticker + "DailyPrice.png", ticker)
        document.add_picture(ticker + "Final.png", width = Inches(8))

    p = document.add_paragraph()
    p.add_run("\n")

    # Insert general ticker info into a table
    table = document.add_table(rows = 1, cols = 4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = keys[0]

    para = hdr_cells[1].paragraphs[0]
    run = para.add_run(values[0])
    run.bold = True

    hdr_cells[2].text = keys[1]

    para = hdr_cells[3].paragraphs[0]
    run = para.add_run(values[1])
    run.bold = True

    for i in range(2, (len(keys) + 1) // 2):
        hdr_cells = table.add_row().cells
        hdr_cells[0].text = keys[i]

        para = hdr_cells[1].paragraphs[0]
        run = para.add_run(values[i])
        run.bold = True

        hdr_cells[2].text = keys[i + 1]

        para = hdr_cells[3].paragraphs[0]
        run = para.add_run(values[i + 1])
        run.bold = True

    return

# Create and format initial document
def createDoc(tickers, T):

    if os.path.exists("Report.docx"):
        os.remove("Report.docx")

    document = Document()

    createCover(document)
    createTOC(document, tickers)

    # Customize document formatting
    sections = document.sections
    for section in sections:
        section.left_margin = Inches(.25)
        section.right_margin = Inches(.25)
        section.top_margin = Inches(.5)
        section.bottom_margin = Inches(.5)

    # Create a new page for each ticker
    for i in range(1, len(tickers) + 1):
        header = document.sections[0].header
        header.is_linked_to_previous = False
        header.text = i + 1
        writeTicker(document, tickers[i - 1], T, i + 1)

        if i < (len(tickers)):
            document.add_page_break()

    document.save("Report.docx")

    return

def main():

    tickers = ["^DJI", "^GSPC", "AAPL", "MSFT", "FB", "GOOG", "INTC",
               "QQQ", "BRK-A", "TSLA"]
    #tickers = ["^DJI", "AAPL"]
    T = StockData.createImages(tickers)

    createDoc(tickers, T)

    files = ["DailyPrice.png", "DailyVolume.png", "Final.png"]
    cleanUpFiles(tickers, files)

    return

main()
